from __future__ import annotations

import argparse
import json
import math
from datetime import datetime
from pathlib import Path
import sys
from typing import Any

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.text import WD_BREAK, WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "src"
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))
if str(SRC) not in sys.path:
    sys.path.insert(0, str(SRC))

from bci_autoresearch.data.vicon_loader import load_vicon_csv
from bci_autoresearch.eval.vicon_ground_plane_report import (
    build_day_statistics_rows,
    choose_window_start_s,
    discover_motion_days,
    select_representative_file_rows,
)


COLOR_BY_SIGNAL = {
    "RHTOE_z": "#4f7cff",
    "RFTOE_z": "#ff5a52",
    "RMTP_z": "#1f9d8b",
    "RANK_z": "#8d63ff",
}
PREFERRED_SIGNALS = ("RHTOE_z", "RFTOE_z", "RMTP_z", "RANK_z")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--motion-base-dir",
        type=Path,
        default=Path("/Volumes/Elements/bci/处理后的关节数据"),
    )
    parser.add_argument(
        "--day-summary-csv",
        type=Path,
        default=Path("/Volumes/Elements/bci/处理后的关节数据/全部日期_marker按天汇总.csv"),
    )
    parser.add_argument(
        "--detail-summary-csv",
        type=Path,
        default=Path("/Volumes/Elements/bci/处理后的关节数据/全部日期_marker完整性统计.csv"),
    )
    parser.add_argument(
        "--figure-output-dir",
        type=Path,
        default=ROOT / "artifacts" / "gait_phase_benchmark" / "vicon_ground_plane_report_all_days",
    )
    parser.add_argument(
        "--report-path",
        type=Path,
        default=ROOT / "output" / "doc" / "步态运动数据按天地面支撑态筛查报告_2026-04-15.docx",
    )
    parser.add_argument("--window-seconds", type=float, default=10.0)
    return parser.parse_args()


def _load_record(path: Path):
    return load_vicon_csv(
        path,
        time_column=None,
        frame_column=None,
        fps=None,
        joints=None,
        target_mode="markers_xyz",
    )


def _extract_signal_map(record: Any) -> dict[str, np.ndarray]:
    name_to_index = {name: idx for idx, name in enumerate(record.names)}
    return {
        signal_name: np.asarray(record.kinematics[:, name_to_index[signal_name]], dtype=np.float32)
        for signal_name in PREFERRED_SIGNALS
        if signal_name in name_to_index
    }


def _lower_envelope_slope(time_s: np.ndarray, values: np.ndarray, *, start_s: float, end_s: float) -> float | None:
    idx = np.nonzero((time_s >= start_s) & (time_s < end_s))[0]
    if idx.size < 100:
        return None
    time_window = time_s[idx]
    value_window = values[idx]
    finite_mask = np.isfinite(value_window)
    if finite_mask.mean() < 0.995:
        return None
    time_window = time_window[finite_mask]
    value_window = value_window[finite_mask]
    if time_window.size < 50:
        return None

    bin_edges = np.arange(start_s, end_s + 1e-9, 1.0, dtype=np.float64)
    floor_times: list[float] = []
    floor_values: list[float] = []
    for left, right in zip(bin_edges[:-1], bin_edges[1:], strict=True):
        mask = (time_window >= left) & (time_window < right)
        if mask.sum() < 20:
            continue
        floor_times.append((left + right) / 2.0)
        floor_values.append(float(np.nanpercentile(value_window[mask], 10)))
    if len(floor_times) < 5:
        return None
    coeff = np.polyfit(np.asarray(floor_times, dtype=np.float64), np.asarray(floor_values, dtype=np.float64), 1)
    return float(coeff[0])


def _segment_q10(values: np.ndarray, *, start_idx: int, end_idx: int) -> float | None:
    segment = np.asarray(values[start_idx:end_idx], dtype=np.float32)
    finite = segment[np.isfinite(segment)]
    if finite.size < 50:
        return None
    return float(np.nanpercentile(finite, 10))


def _build_day_figure(
    *,
    day: str,
    motion_base_dir: Path,
    selected_rows: list[dict[str, Any]],
    output_path: Path,
    window_seconds: float,
) -> list[dict[str, Any]]:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    plt.rcParams["font.family"] = "Arial Unicode MS"
    fig, axes = plt.subplots(max(1, len(selected_rows)), 1, figsize=(12.5, 11.8), squeeze=False)
    fig.patch.set_facecolor("white")
    segment_rows: list[dict[str, Any]] = []

    for ax, row in zip(axes[:, 0], selected_rows, strict=False):
        file_path = motion_base_dir / day / str(row["file"])
        record = _load_record(file_path)
        time_s = np.asarray(record.time_s, dtype=np.float64)
        signal_map = _extract_signal_map(record)
        start_s = choose_window_start_s(
            duration_s=float(row.get("duration_s") or time_s[-1]),
            complete_prefix_s=row.get("complete_prefix_s"),
            window_s=window_seconds,
        )
        end_s = min(float(time_s[-1]), start_s + window_seconds)
        mask = (time_s >= start_s) & (time_s < end_s)
        idx = np.nonzero(mask)[0]
        if idx.size == 0:
            raise ValueError(f"Empty segment for {file_path.name} {start_s}-{end_s}s")
        start_idx = int(idx[0])
        end_idx = int(idx[-1]) + 1

        toe_slopes: list[float] = []
        toe_slope_text: list[str] = []
        for signal_name in PREFERRED_SIGNALS:
            if signal_name not in signal_map:
                continue
            series = signal_map[signal_name]
            ax.plot(
                time_s[start_idx:end_idx],
                series[start_idx:end_idx],
                lw=1.5,
                color=COLOR_BY_SIGNAL[signal_name],
                label=signal_name,
            )
            q10 = _segment_q10(series, start_idx=start_idx, end_idx=end_idx)
            if q10 is not None:
                ax.axhline(q10, color=COLOR_BY_SIGNAL[signal_name], lw=0.9, ls="--", alpha=0.2)
            if signal_name in ("RHTOE_z", "RFTOE_z"):
                slope = _lower_envelope_slope(time_s, series.astype(np.float64), start_s=start_s, end_s=end_s)
                if slope is not None:
                    toe_slopes.append(abs(slope))
                    toe_slope_text.append(f"{signal_name} 低位斜率 {slope:.3f} mm/s")

        segment_rows.append(
            {
                "date": day,
                "file": str(row["file"]),
                "start_s": round(start_s, 3),
                "end_s": round(end_s, 3),
                "abs_toe_slope_mm_per_s": round(float(np.median(toe_slopes)), 4) if toe_slopes else None,
            }
        )

        ax.set_title(
            f"{row['file']} {start_s:.1f}-{end_s:.1f} 秒",
            loc="left",
            fontsize=12,
            color="#17324a",
        )
        if toe_slope_text:
            ax.text(
                0.0,
                1.01,
                " | ".join(toe_slope_text),
                transform=ax.transAxes,
                fontsize=9,
                color="#516575",
            )
        ax.set_ylabel("z 轴位置", fontsize=10, color="#17324a")
        ax.grid(True, axis="y", alpha=0.18)
        ax.spines["top"].set_visible(False)
        ax.spines["right"].set_visible(False)
        ax.tick_params(labelsize=9, colors="#44596d")
        ax.legend(loc="upper right", fontsize=8, frameon=False, ncol=4)

    for ax in axes[:, 0]:
        ax.set_xlabel("时间（秒）", fontsize=10, color="#17324a")

    fig.suptitle(f"{day}: 代表性 3 个试次 x 各 10 秒脚部 z 轴", x=0.02, ha="left", fontsize=16, color="#17324a")
    fig.tight_layout(rect=(0, 0, 1, 0.97))
    fig.savefig(output_path, dpi=180, facecolor=fig.get_facecolor(), bbox_inches="tight")
    plt.close(fig)
    return segment_rows


def _set_document_defaults(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Arial Unicode MS"
    normal.font.size = Pt(10.5)
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)


def _render_report_docx(
    *,
    report_path: Path,
    table_rows: list[dict[str, Any]],
    day_entries: list[dict[str, Any]],
) -> None:
    report_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    _set_document_defaults(doc)

    title = doc.add_heading("步态运动数据按天地面支撑态筛查报告", level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.add_run(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}")

    intro = doc.add_paragraph()
    intro.add_run("口径：").bold = True
    intro.add_run(
        "每个日期选 3 个代表性试次；每个试次截 10 秒；统一绘制 RHTOE_z、RFTOE_z、RMTP_z、RANK_z。"
        " 统计表结合按天 marker 完整性与示例段脚趾低位斜率，用来快速判断哪一天更像有清楚支撑态和平地参考。"
    )

    doc.add_heading("按天统计表", level=1)
    headers = [
        "日期",
        "试次数",
        "完整试次",
        "部分掉点",
        "完整占比",
        "总时长",
        "最早掉点(s)",
        "示例段脚趾低位斜率中位数(mm/s)",
    ]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for cell, header in zip(table.rows[0].cells, headers, strict=True):
        cell.text = header
    for row in table_rows:
        cells = table.add_row().cells
        cells[0].text = str(row["date"])
        cells[1].text = str(row["files"])
        cells[2].text = str(row["all_12_complete_files"])
        cells[3].text = str(row["partial_files"])
        cells[4].text = f"{float(row['complete_file_ratio']):.2%}"
        cells[5].text = str(row["duration_hms_sum"])
        cells[6].text = "-" if row["earliest_dropout_s"] is None else f"{float(row['earliest_dropout_s']):.2f}"
        cells[7].text = "-" if row["median_abs_toe_slope_mm_per_s"] is None else f"{float(row['median_abs_toe_slope_mm_per_s']):.4f}"

    for index, entry in enumerate(day_entries):
        doc.add_heading(entry["date"], level=1)
        stats = entry["stats"]
        para = doc.add_paragraph()
        para.add_run("摘要：").bold = True
        para.add_run(
            f"完整试次 {stats['all_12_complete_files']}/{stats['files']}，"
            f"部分掉点 {stats['partial_files']}，"
            f"最早掉点 "
            f"{'-' if stats['earliest_dropout_s'] is None else f'{float(stats['earliest_dropout_s']):.2f} 秒'}，"
            f"示例段脚趾低位斜率中位数 "
            f"{'-' if stats['median_abs_toe_slope_mm_per_s'] is None else f'{float(stats['median_abs_toe_slope_mm_per_s']):.4f} mm/s'}。"
        )
        files_para = doc.add_paragraph()
        files_para.add_run("示例试次：").bold = True
        files_para.add_run(" / ".join(seg["file"] for seg in entry["segments"]))
        doc.add_picture(str(entry["figure_path"]), width=Inches(6.8))
        if index != len(day_entries) - 1:
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    doc.save(report_path)


def main() -> None:
    args = parse_args()
    day_summary_df = pd.read_csv(args.day_summary_csv, dtype={"date": str})
    detail_df = pd.read_csv(args.detail_summary_csv, dtype={"date": str})

    days = discover_motion_days(args.motion_base_dir)
    figure_dir = args.figure_output_dir
    figure_dir.mkdir(parents=True, exist_ok=True)

    all_segment_rows: list[dict[str, Any]] = []
    day_entries: list[dict[str, Any]] = []

    for day in days:
        detail_rows = detail_df[detail_df["date"] == day].copy()
        if detail_rows.empty:
            continue
        selected_rows = select_representative_file_rows(detail_rows.to_dict("records"), top_k=3)
        if not selected_rows:
            continue
        figure_path = figure_dir / f"{day}_representative3_ground_plane_check.png"
        segment_rows = _build_day_figure(
            day=day,
            motion_base_dir=args.motion_base_dir,
            selected_rows=selected_rows,
            output_path=figure_path,
            window_seconds=float(args.window_seconds),
        )
        all_segment_rows.extend(segment_rows)
        day_entries.append(
            {
                "date": day,
                "figure_path": figure_path,
                "segments": segment_rows,
            }
        )

    table_rows = build_day_statistics_rows(day_summary_df.to_dict("records"), all_segment_rows)
    stats_by_day = {row["date"]: row for row in table_rows}
    for entry in day_entries:
        entry["stats"] = stats_by_day[entry["date"]]

    _render_report_docx(
        report_path=args.report_path,
        table_rows=table_rows,
        day_entries=day_entries,
    )

    metadata = {
        "report_path": str(args.report_path),
        "figure_dir": str(figure_dir),
        "days": [entry["date"] for entry in day_entries],
        "table_rows": table_rows,
        "segments": all_segment_rows,
    }
    (figure_dir / "report_metadata.json").write_text(json.dumps(metadata, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps({"report_path": str(args.report_path), "figure_count": len(day_entries)}, ensure_ascii=False))


if __name__ == "__main__":
    main()
