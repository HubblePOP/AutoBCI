from __future__ import annotations

import argparse
import json
import math
from datetime import datetime
from pathlib import Path
import sys
from typing import Any

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.text import WD_BREAK, WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "src"
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))
if str(SRC) not in sys.path:
    sys.path.insert(0, str(SRC))

from bci_autoresearch.data.vicon_loader import load_vicon_csv
from bci_autoresearch.eval.gait_phase import compute_hysteresis_reference_trace
from bci_autoresearch.eval.vicon_ground_plane_head_mid_tail import (
    build_head_mid_tail_windows,
    summarize_sampled_swing_ratios,
    summarize_ground_visibility,
)
from bci_autoresearch.eval.vicon_ground_plane_report import (
    build_day_statistics_rows,
    discover_motion_days,
    select_representative_file_rows,
)


SIGNAL_ORDER = ("RHTOE_z", "RFTOE_z", "RMTP_z", "RANK_z")
SIGNAL_COLORS = {
    "RHTOE_z": "#3b82f6",
    "RFTOE_z": "#ef4444",
    "RMTP_z": "#14b8a6",
    "RANK_z": "#8b5cf6",
}


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Build per-day head/mid/tail ground-plane report DOCX.")
    parser.add_argument(
        "--motion-base-dir",
        type=Path,
        default=Path("/Volumes/Elements/bci/处理后的关节数据"),
    )
    parser.add_argument(
        "--day-summary-csv",
        type=Path,
        default=Path("/Volumes/Elements/bci/处理后的关节数据/全部日期_marker按天汇总.csv"),
    )
    parser.add_argument(
        "--detail-summary-csv",
        type=Path,
        default=Path("/Volumes/Elements/bci/处理后的关节数据/全部日期_marker完整性统计.csv"),
    )
    parser.add_argument(
        "--figure-output-dir",
        type=Path,
        default=ROOT / "artifacts" / "gait_phase_benchmark" / "vicon_ground_plane_head_mid_tail_by_day",
    )
    parser.add_argument(
        "--report-path",
        type=Path,
        default=ROOT / "output" / "doc" / "步态运动数据按日期前中后10秒地面平面抽样报告_2026-04-16.docx",
    )
    parser.add_argument("--window-seconds", type=float, default=10.0)
    return parser.parse_args()


def _load_record(path: Path):
    return load_vicon_csv(
        path,
        time_column=None,
        frame_column=None,
        fps=None,
        joints=None,
        target_mode="markers_xyz",
    )


def _extract_signal_map(record: Any) -> dict[str, np.ndarray]:
    name_to_index = {name: idx for idx, name in enumerate(record.names)}
    return {
        signal_name: np.asarray(record.kinematics[:, name_to_index[signal_name]], dtype=np.float32)
        for signal_name in SIGNAL_ORDER
        if signal_name in name_to_index
    }


def _lower_envelope_slope(time_s: np.ndarray, values: np.ndarray, *, start_s: float, end_s: float) -> float | None:
    idx = np.nonzero((time_s >= start_s) & (time_s < end_s))[0]
    if idx.size < 100:
        return None

    time_window = time_s[idx]
    value_window = values[idx]
    finite_mask = np.isfinite(value_window)
    if finite_mask.mean() < 0.99:
        return None

    time_window = time_window[finite_mask]
    value_window = value_window[finite_mask]
    if time_window.size < 50:
        return None

    bin_edges = np.arange(start_s, end_s + 1e-9, 1.0, dtype=np.float64)
    floor_times: list[float] = []
    floor_values: list[float] = []
    for left, right in zip(bin_edges[:-1], bin_edges[1:], strict=True):
        mask = (time_window >= left) & (time_window < right)
        if mask.sum() < 20:
            continue
        floor_times.append((left + right) / 2.0)
        floor_values.append(float(np.nanpercentile(value_window[mask], 10)))

    if len(floor_times) < 5:
        return None

    coeff = np.polyfit(
        np.asarray(floor_times, dtype=np.float64),
        np.asarray(floor_values, dtype=np.float64),
        1,
    )
    return float(coeff[0])


def _segment_quantile(values: np.ndarray, *, start_idx: int, end_idx: int, q: float = 10.0) -> float | None:
    segment = np.asarray(values[start_idx:end_idx], dtype=np.float32)
    finite = segment[np.isfinite(segment)]
    if finite.size < 50:
        return None
    return float(np.nanpercentile(finite, q))


def _mask_from_intervals(length: int, intervals: list[dict[str, int]]) -> np.ndarray:
    mask = np.zeros(length, dtype=np.uint8)
    for item in intervals:
        start_idx = int(item["start_idx"])
        end_idx = int(item["end_idx"])
        if end_idx > start_idx:
            mask[start_idx:end_idx] = 1
    return mask


def _set_document_defaults(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Arial Unicode MS"
    normal.font.size = Pt(10.5)
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)


def _build_day_figure(
    *,
    day: str,
    selected_row: dict[str, Any],
    motion_base_dir: Path,
    output_path: Path,
    window_seconds: float,
) -> dict[str, Any]:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    file_name = str(selected_row["file"])
    file_path = motion_base_dir / day / file_name
    record = _load_record(file_path)
    time_s = np.asarray(record.time_s, dtype=np.float64)
    signal_map = _extract_signal_map(record)
    duration_s = float(time_s[-1]) if time_s.size else 0.0
    windows = build_head_mid_tail_windows(duration_s=duration_s, window_s=window_seconds)
    trace_map = {
        signal_name: compute_hysteresis_reference_trace(
            time_s=time_s,
            toe_z=signal_map[signal_name],
            signal_name=signal_name,
        )
        for signal_name in ("RHTOE_z", "RFTOE_z")
        if signal_name in signal_map
    }
    swing_mask_map = {
        signal_name: _mask_from_intervals(time_s.shape[0], list(trace["swing_intervals"]))
        for signal_name, trace in trace_map.items()
    }

    plt.rcParams["font.family"] = "Arial Unicode MS"
    fig, axes = plt.subplots(3, 1, figsize=(12.8, 9.6), squeeze=False, sharey=False)
    fig.patch.set_facecolor("white")

    window_rows: list[dict[str, Any]] = []
    for ax, (window_label, start_s, end_s) in zip(axes[:, 0], windows, strict=True):
        idx = np.nonzero((time_s >= start_s) & (time_s <= end_s))[0]
        if idx.size == 0:
            raise ValueError(f"Empty segment for {file_name} {start_s}-{end_s}s")
        start_idx = int(idx[0])
        end_idx = int(idx[-1]) + 1
        x = time_s[start_idx:end_idx]

        toe_slopes: list[float] = []
        toe_slope_labels: list[str] = []
        for signal_name in SIGNAL_ORDER:
            if signal_name not in signal_map:
                continue
            series = signal_map[signal_name]
            y = series[start_idx:end_idx]
            ax.plot(
                x,
                y,
                lw=1.25,
                color=SIGNAL_COLORS[signal_name],
                label=signal_name,
                alpha=0.94,
            )
            if signal_name in ("RHTOE_z", "RFTOE_z"):
                q10 = _segment_quantile(series, start_idx=start_idx, end_idx=end_idx, q=10.0)
                if q10 is not None:
                    ax.axhline(q10, color=SIGNAL_COLORS[signal_name], lw=0.9, ls="--", alpha=0.22)
                slope = _lower_envelope_slope(time_s, series.astype(np.float64), start_s=start_s, end_s=end_s)
                if slope is not None:
                    toe_slopes.append(abs(slope))
                    toe_slope_labels.append(f"{signal_name} 低位斜率 {slope:.3f} mm/s")

        median_abs_slope = float(np.median(toe_slopes)) if toe_slopes else None
        window_rows.append(
            {
                "date": day,
                "file": file_name,
                "window_label": window_label,
                "start_s": round(start_s, 3),
                "end_s": round(end_s, 3),
                "abs_toe_slope_mm_per_s": round(median_abs_slope, 4) if median_abs_slope is not None else None,
                "rh_swing_ratio": round(float(swing_mask_map["RHTOE_z"][start_idx:end_idx].mean()), 4)
                if "RHTOE_z" in swing_mask_map
                else None,
                "rf_swing_ratio": round(float(swing_mask_map["RFTOE_z"][start_idx:end_idx].mean()), 4)
                if "RFTOE_z" in swing_mask_map
                else None,
            }
        )

        ax.set_title(f"{window_label} · {start_s:.1f}-{end_s:.1f} 秒", loc="left", fontsize=11.5, color="#17324a")
        if toe_slope_labels:
            ax.text(
                1.0,
                1.02,
                " | ".join(toe_slope_labels),
                transform=ax.transAxes,
                fontsize=8.8,
                color="#516575",
                ha="right",
            )
        ax.set_ylabel("z 轴位置", fontsize=9.8, color="#17324a")
        ax.grid(True, axis="y", alpha=0.18)
        ax.spines["top"].set_visible(False)
        ax.spines["right"].set_visible(False)
        ax.tick_params(labelsize=9, colors="#44596d")

    axes[-1, 0].set_xlabel("时间（秒）", fontsize=10, color="#17324a")
    axes[0, 0].legend(loc="upper right", fontsize=8, frameon=False, ncol=4)

    fig.suptitle(
        f"{day} · {file_name} · 同一代表试次的前 / 中 / 后各 10 秒脚部 z 轴",
        x=0.02,
        ha="left",
        fontsize=15.5,
        color="#17324a",
    )
    fig.tight_layout(rect=(0, 0, 1, 0.97))
    fig.savefig(output_path, dpi=180, facecolor=fig.get_facecolor(), bbox_inches="tight")
    plt.close(fig)

    return {
        "date": day,
        "file": file_name,
        "file_path": str(file_path),
        "figure_path": str(output_path),
        "duration_s": round(duration_s, 3),
        "complete_prefix_s": selected_row.get("complete_prefix_s"),
        "complete_markers": int(selected_row.get("complete_markers") or 0),
        "incomplete_markers": int(selected_row.get("incomplete_markers") or 0),
        "windows": window_rows,
        **summarize_sampled_swing_ratios(window_rows),
    }


def _render_report_docx(
    *,
    report_path: Path,
    table_rows: list[dict[str, Any]],
    day_entries: list[dict[str, Any]],
) -> None:
    report_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    _set_document_defaults(doc)

    title = doc.add_heading("按日期前中后 10 秒抽样的步态运动地面平面检查", level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    meta = doc.add_paragraph()
    meta.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    meta.add_run(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}")

    intro = doc.add_paragraph()
    intro.add_run("说明：").bold = True
    intro.add_run(
        "每个日期只选一个代表试次。对同一试次取开头、中间、结尾各 10 秒，统一画出 "
        "RHTOE_z、RFTOE_z、RMTP_z、RANK_z，用来肉眼判断该日期是否较容易看出稳定的地面支撑平面。"
    )

    note = doc.add_paragraph()
    note.add_run("判读提示：").bold = True
    note.add_run(
        "如果脚趾 z 轴在低位时能形成近似水平的平台，通常更容易把支撑态看出来。"
        " 本报告里的“平面可见性”只是一个初筛标签，最终仍以图像本身为准。"
    )

    doc.add_heading("按天统计表", level=1)
    headers = [
        "日期",
        "代表试次",
        "后肢摆动占比",
        "前肢摆动占比",
        "完整试次占比",
        "最早掉点(s)",
        "三段脚趾低位斜率中位数(mm/s)",
        "平面可见性",
    ]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for cell, header in zip(table.rows[0].cells, headers, strict=True):
        cell.text = header

    for row in table_rows:
        cells = table.add_row().cells
        cells[0].text = str(row["date"])
        cells[1].text = str(row.get("selected_file") or "-")
        cells[2].text = "-" if row.get("rh_swing_ratio_median") is None else f"{float(row['rh_swing_ratio_median']):.2%}"
        cells[3].text = "-" if row.get("rf_swing_ratio_median") is None else f"{float(row['rf_swing_ratio_median']):.2%}"
        cells[4].text = f"{float(row['complete_file_ratio']):.2%}"
        cells[5].text = "-" if row["earliest_dropout_s"] is None else f"{float(row['earliest_dropout_s']):.2f}"
        cells[6].text = "-" if row["median_abs_toe_slope_mm_per_s"] is None else f"{float(row['median_abs_toe_slope_mm_per_s']):.4f}"
        cells[7].text = str(row.get("ground_visibility") or "-")

    for index, entry in enumerate(day_entries):
        doc.add_heading(str(entry["date"]), level=1)
        para = doc.add_paragraph()
        para.add_run("代表试次：").bold = True
        para.add_run(str(entry["file"]))
        para.add_run("；").bold = False
        para.add_run("总时长：").bold = True
        para.add_run(f"{float(entry['duration_s']):.1f} 秒")
        para.add_run("；").bold = False
        para.add_run("平面可见性：").bold = True
        para.add_run(str(entry["ground_visibility"]))

        swing_para = doc.add_paragraph()
        swing_para.add_run("抽样摆动占比：").bold = True
        swing_para.add_run(
            "后肢 "
            + ("-" if entry.get("rh_swing_ratio_median") is None else f"{float(entry['rh_swing_ratio_median']):.2%}")
            + "；前肢 "
            + ("-" if entry.get("rf_swing_ratio_median") is None else f"{float(entry['rf_swing_ratio_median']):.2%}")
        )

        slope_para = doc.add_paragraph()
        slope_para.add_run("三段脚趾低位斜率中位数：").bold = True
        slope_value = entry.get("median_abs_toe_slope_mm_per_s")
        slope_para.add_run("-" if slope_value is None else f"{float(slope_value):.4f} mm/s")

        range_para = doc.add_paragraph()
        range_para.add_run("抽样区段：").bold = True
        range_para.add_run(
            " / ".join(
                f"{window['window_label']} {float(window['start_s']):.1f}-{float(window['end_s']):.1f} 秒"
                for window in entry["windows"]
            )
        )

        doc.add_picture(str(entry["figure_path"]), width=Inches(6.9))
        if index != len(day_entries) - 1:
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    doc.save(report_path)


def main() -> None:
    args = parse_args()
    day_summary_df = pd.read_csv(args.day_summary_csv, dtype={"date": str})
    detail_df = pd.read_csv(args.detail_summary_csv, dtype={"date": str})

    figure_dir = args.figure_output_dir
    figure_dir.mkdir(parents=True, exist_ok=True)
    days = discover_motion_days(args.motion_base_dir)

    all_segment_rows: list[dict[str, Any]] = []
    day_entries: list[dict[str, Any]] = []

    for day in days:
        detail_rows = detail_df[detail_df["date"] == day].copy()
        if detail_rows.empty:
            continue
        selected_rows = select_representative_file_rows(detail_rows.to_dict("records"), top_k=1)
        if not selected_rows:
            continue
        selected_row = selected_rows[0]
        figure_path = figure_dir / f"{day}_head_mid_tail_ground_plane.png"
        entry = _build_day_figure(
            day=day,
            selected_row=selected_row,
            motion_base_dir=args.motion_base_dir,
            output_path=figure_path,
            window_seconds=float(args.window_seconds),
        )
        all_segment_rows.extend(entry["windows"])
        day_entries.append(entry)

    table_rows = build_day_statistics_rows(day_summary_df.to_dict("records"), all_segment_rows)
    stats_by_day = {row["date"]: row for row in table_rows}
    entry_by_day = {entry["date"]: entry for entry in day_entries}

    enriched_rows: list[dict[str, Any]] = []
    for day in days:
        if day not in stats_by_day or day not in entry_by_day:
            continue
        row = dict(stats_by_day[day])
        entry = entry_by_day[day]
        row["selected_file"] = entry["file"]
        row["ground_visibility"] = summarize_ground_visibility(
            row.get("median_abs_toe_slope_mm_per_s"),
            row.get("complete_file_ratio", 0.0),
        )
        row["rh_swing_ratio_median"] = entry.get("rh_swing_ratio_median")
        row["rf_swing_ratio_median"] = entry.get("rf_swing_ratio_median")
        enriched_rows.append(row)

        entry["median_abs_toe_slope_mm_per_s"] = row.get("median_abs_toe_slope_mm_per_s")
        entry["ground_visibility"] = row["ground_visibility"]

    _render_report_docx(
        report_path=args.report_path,
        table_rows=enriched_rows,
        day_entries=[entry_by_day[row["date"]] for row in enriched_rows],
    )

    metadata = {
        "report_path": str(args.report_path),
        "figure_dir": str(figure_dir),
        "days": [row["date"] for row in enriched_rows],
        "table_rows": enriched_rows,
        "day_entries": day_entries,
    }
    (figure_dir / "report_metadata.json").write_text(
        json.dumps(metadata, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    print(
        json.dumps(
            {
                "report_path": str(args.report_path),
                "days": len(enriched_rows),
                "figure_dir": str(figure_dir),
            },
            ensure_ascii=False,
        )
    )


if __name__ == "__main__":
    main()
